"""
All HTML parsing logic. Each fetch_* function takes raw HTML (str) so it can
be unit-tested with fixture files without hitting the network.
The fetch_*_live wrappers handle HTTP + cache.
"""
import re
from typing import Any
from bs4 import BeautifulSoup, Tag

from . import cache
from .auth import authed_get, get_client, login
from . import config
from .context import get_current_user


def _base() -> str:
    """Current user's ManageBac base URL (per-user; schools differ)."""
    u = get_current_user()
    return u.mb_url if u else config.BASE_URL


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _text(el: Tag | None) -> str:
    return el.get_text(strip=True) if el else ""


def _html_to_markdown(el) -> str:
    """
    Convert a BeautifulSoup element to Markdown, preserving rich formatting.

    ManageBac uses inline styles rather than semantic tags:
      style="font-weight: bold"    → **bold**
      style="font-style: italic"   → *italic*
      style="text-decoration: underline" → __underline__ (shown as bold too)
    Standard tags (strong, em, h1-h6, ul, ol, a, br) are also handled.
    """
    from bs4 import NavigableString

    def node(el) -> str:
        if isinstance(el, NavigableString):
            return str(el)

        tag = getattr(el, "name", None)
        if tag is None:
            return ""

        style = el.get("style", "")
        bold = ("font-weight: bold" in style or "font-weight:bold" in style
                or tag in ("strong", "b"))
        italic = ("font-style: italic" in style or "font-style:italic" in style
                  or tag in ("em", "i"))

        def kids() -> str:
            return "".join(node(c) for c in el.children)

        # --- block elements ---
        if tag == "br":
            return "\n"
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            return "#" * int(tag[1]) + " " + kids().strip() + "\n\n"
        if tag == "p":
            t = kids().strip()
            return (t + "\n\n") if t else ""
        if tag in ("div", "section"):
            t = kids().strip()
            return (t + "\n\n") if t else ""
        if tag == "ul":
            items = ["- " + "".join(node(c) for c in li.children).strip()
                     for li in el.find_all("li", recursive=False)]
            return ("\n".join(items) + "\n\n") if items else ""
        if tag == "ol":
            items = [f"{i}. " + "".join(node(c) for c in li.children).strip()
                     for i, li in enumerate(el.find_all("li", recursive=False), 1)]
            return ("\n".join(items) + "\n\n") if items else ""
        if tag == "li":
            return "- " + kids().strip() + "\n"

        # --- inline elements ---
        if tag == "a":
            href = el.get("href", "")
            a_classes = el.get("class", [])
            # Embedded file attachment — render as a plain label, not a link
            # (the URL is captured separately in embedded_files)
            if "fr-file" in a_classes or el.get("data-name"):
                name = el.get("data-name") or ""
                size_el = el.find(class_="fr-file-size")
                size = size_el.get_text(strip=True) if size_el else ""
                label = name or el.get_text(strip=True)
                return f"📎 {label}" + (f" ({size})" if size else "")
            text = kids().strip()
            if href.startswith("http"):
                label = text if (text and text != href) else href
                return f"[{label}]({href})"
            return kids()

        if tag in ("span", "strong", "b", "em", "i", "u"):
            text = kids()
            stripped = text.strip()
            if not stripped:
                return text
            if bold and italic:
                return text.replace(stripped, f"***{stripped}***", 1)
            elif bold:
                return text.replace(stripped, f"**{stripped}**", 1)
            elif italic:
                return text.replace(stripped, f"*{stripped}*", 1)
            return text

        # everything else — just recurse
        return kids()

    md = node(el)
    md = re.sub(r'\n{3,}', '\n\n', md)   # collapse 3+ blank lines → 2
    md = re.sub(r' +\n', '\n', md)        # strip trailing spaces
    return md.strip()


def _parse_grade_box(cell: Tag) -> dict[str, dict]:
    """Parse grade cells like B: [7|8] C: [5|8] into {B: {score:7, max:8}}"""
    grades = {}
    # Each criterion label + score box pair
    labels = cell.find_all(string=re.compile(r'^[A-D]:$'))
    for label in labels:
        label_text = label.strip().rstrip(":")
        # Score boxes are siblings: two consecutive elements with numbers
        boxes = []
        sibling = label.parent.next_sibling if label.parent else None
        # Walk siblings looking for score boxes
        for _ in range(6):
            if sibling is None:
                break
            if hasattr(sibling, 'get_text'):
                t = sibling.get_text(strip=True)
                if re.match(r'^\d+$', t):
                    boxes.append(int(t))
                elif t and t not in (':', '|'):
                    if boxes:
                        break
            sibling = getattr(sibling, 'next_sibling', None)
        if len(boxes) >= 2:
            grades[label_text] = {"score": boxes[0], "max": boxes[1]}
        elif len(boxes) == 1:
            grades[label_text] = {"score": boxes[0], "max": None}
    return grades


def _parse_grades_from_task_row(row: Tag) -> dict[str, dict]:
    """
    Parse grades from the right-side grade column of a task row.
    Handles formats like: A: [7|8]  or  B: [7|8] C: [5|8]  or  N/A
    """
    grades = {}
    # Find the grade column — it's the last cell / div in the row
    grade_col = row.find(class_=re.compile(r'grade|criterion|score', re.I))
    if not grade_col:
        # Try to find by position — look for elements with number boxes
        all_divs = row.find_all("div")
        for div in reversed(all_divs):
            text = div.get_text()
            if re.search(r'[A-D]:\s*\d', text):
                grade_col = div
                break

    if not grade_col:
        return grades

    text = grade_col.get_text(" ", strip=True)
    # Pattern: "A: 7 8" or "B: 7 8 C: 5 8"
    for match in re.finditer(r'([A-D]):\s*(\d+)\s+(\d+)', text):
        criterion, score, max_score = match.groups()
        grades[criterion] = {"score": int(score), "max": int(max_score)}
    # Also capture N/A grades: "B: N/A" or "B:N/A"
    for match in re.finditer(r'\b([A-D]):\s*N/A\b', text, re.IGNORECASE):
        criterion = match.group(1).upper()
        if criterion not in grades:
            grades[criterion] = {"score": None, "max": None}

    return grades


# ---------------------------------------------------------------------------
# Classes
# ---------------------------------------------------------------------------

def parse_classes(html: str) -> list[dict]:
    soup = BeautifulSoup(html, "lxml")
    classes = []

    # Each class entry is a link in the left sidebar or class list
    # URL pattern: /student/classes/{id}
    for a in soup.find_all("a", href=re.compile(r'/student/classes/\d+$')):
        href = a["href"]
        class_id = re.search(r'/classes/(\d+)', href).group(1)
        name = a.get_text(strip=True)
        if not name or name in ("Browse All Classes", "Classes"):
            continue

        # Detect level tags (HL/SL) near this element
        parent = a.parent
        tags = [t.get_text(strip=True) for t in parent.find_all(
            class_=re.compile(r'tag|label|badge', re.I))] if parent else []

        classes.append({
            "id": class_id,
            "name": name,
            "url": f"{_base()}/student/classes/{class_id}",
            "level_tags": tags,
        })

    # Deduplicate by id (sidebar + main list may repeat)
    seen = set()
    unique = []
    for c in classes:
        if c["id"] not in seen:
            seen.add(c["id"])
            unique.append(c)
    return unique


async def fetch_classes() -> list[dict]:
    # `if cached:` (not `is not None`) deliberately treats a cached EMPTY list as
    # a miss. A student always has classes, so an empty result only ever means a
    # transient failure (login page, ManageBac hiccup). Caching that empty list
    # for 24h is what poisoned get_classes — and every tool that depends on it
    # (get_grades, tag_search, get_upcoming) — for a whole day.
    cached = cache.get("get_classes")
    if cached:
        return cached

    import asyncio as _asyncio

    async with await get_client() as client:
        # authed_get already re-logs-in transparently on a /login redirect.
        r = await authed_get(client, "/student/classes/my")
        result = parse_classes(r.text)

        # Never cache an empty parse — return it so the NEXT call retries fresh
        # instead of serving stale-empty until the TTL expires.
        if not result:
            return []

        # Check each class's journal tab concurrently (the request semaphore in
        # auth.py caps real parallelism). Sequentially this was ~18 round-trips,
        # slow enough to hit the connector's timeout.
        pages = await _asyncio.gather(
            *[authed_get(client, f"/student/classes/{c['id']}") for c in result]
        )

    for cls, r2 in zip(result, pages):
        cls["has_journal"] = "learner_portfolio" in r2.text

    cache.set("get_classes", result, "get_classes")
    return result


# ---------------------------------------------------------------------------
# Timetable
# ---------------------------------------------------------------------------

def parse_timetable(html: str) -> list[dict]:
    """
    Parse the weekly timetable from table.f-timetable.

    Structure:
      thead > tr > th[0]="Period" th[1]="Jun 1, Mon" ... th[5]="Jun 5, Fri"
      tbody > tr (one per period):
        th > span.f-numeric      → period number  (or small text for "Homeroom")
        td (one per day):
          a.f-timetable-item     → one class per cell (sometimes multiple)
            div.f-box-item__body:
              small.color-secondary    → "9:00 AM - 9:45 AM"
              div.badge.tasks-counter  → optional task count badge
                span.badge-label       → "1"
              p.fw-semibold.text-truncate → class name
              p.text-truncate.mt-1    → grade (skip)
              p.text-truncate         → teacher name
              p.text-truncate         → room number (optional)
            data-bs-content-url has ib_class_id=12734184 for class_id
    """
    soup = BeautifulSoup(html, "lxml")
    slots = []

    table = soup.find("table", class_="f-timetable")
    if not table:
        return slots

    # Day names from header row
    thead = table.find("thead")
    days = []
    if thead:
        for th in thead.find_all("th")[1:]:  # skip "Period" column
            days.append(th.get_text(strip=True))

    tbody = table.find("tbody")
    if not tbody:
        return slots

    for row in tbody.find_all("tr"):
        # Period number — span.f-numeric for numbered periods, small for named (Homeroom)
        row_th = row.find("th")
        if not row_th:
            continue
        numeric_el = row_th.find("span", class_="f-numeric")
        if numeric_el:
            period_text = numeric_el.get_text(strip=True)
            try:
                period = int(period_text)
            except ValueError:
                continue
            period_name = str(period)
        else:
            period_name = row_th.get_text(strip=True)
            period = 0  # Homeroom / special rows

        for i, cell in enumerate(row.find_all("td")):
            day = days[i] if i < len(days) else f"Day{i+1}"

            # Each class in this cell is an a.f-timetable-item
            for item in cell.find_all("a", class_="f-timetable-item"):
                body = item.find("div", class_="f-box-item__body")
                if not body:
                    continue

                # Time
                time_el = body.find("small", class_="color-secondary")
                time_text = time_el.get_text(strip=True) if time_el else ""
                times = re.findall(r'\d+:\d+\s*[AP]M', time_text)

                # Task count
                badge_label = body.find("span", class_="badge-label")
                task_count = 0
                if badge_label:
                    try:
                        task_count = int(badge_label.get_text(strip=True))
                    except ValueError:
                        pass

                # Class name, teacher, room — the <p> elements in order:
                # p.fw-semibold = class name, p.mt-1 = grade (skip), remaining p = teacher, room
                paragraphs = body.find_all("p")
                class_name = ""
                teacher = ""
                room = ""
                content_ps = []
                for p in paragraphs:
                    p_classes = p.get("class", [])
                    if "fw-semibold" in p_classes:
                        class_name = p.get_text(strip=True)
                    elif "mt-1" not in p_classes:
                        content_ps.append(p.get_text(strip=True))
                if content_ps:
                    teacher = content_ps[0]
                if len(content_ps) > 1:
                    room = content_ps[1]

                # Class ID from URL param
                content_url = item.get("data-bs-content-url", "")
                class_id_match = re.search(r'ib_class_id=(\d+)', content_url)
                class_id = class_id_match.group(1) if class_id_match else ""

                if class_name:
                    slots.append({
                        "period": period if period else period_name,
                        "day": day,
                        "time_start": times[0] if times else "",
                        "time_end": times[1] if len(times) > 1 else "",
                        "class_name": class_name,
                        "class_id": class_id,
                        "teacher": teacher,
                        "room": room,
                        "task_count": task_count,
                    })

    return slots


def _now_info() -> dict:
    """Current date/time in the server's (school's) local timezone."""
    import datetime
    now = datetime.datetime.now().astimezone()
    return {
        "weekday": now.strftime("%A"),
        "date": now.strftime("%B ") + str(now.day) + now.strftime(", %Y"),
        "time": now.strftime("%I:%M %p").lstrip("0"),
        "iso": now.isoformat(timespec="minutes"),
        "timezone": now.strftime("%Z") or str(now.utcoffset()),
    }


async def fetch_timetable() -> dict:
    # Timetable slots are cached (they change rarely); the "current" time is
    # always computed fresh so the AI knows what day/time it actually is.
    # `if slots:` treats a cached EMPTY timetable as a miss — a student always
    # has a timetable, so empty means a transient fetch failure, not real data.
    slots = cache.get("get_timetable")
    if not slots:
        async with await get_client() as client:
            r = await authed_get(client, "/student/timetables")
        slots = parse_timetable(r.text)
        if slots:                       # never cache an empty parse
            cache.set("get_timetable", slots, "get_timetable")

    return {"current": _now_info(), "timetable": slots}


# ---------------------------------------------------------------------------
# Tasks list
# ---------------------------------------------------------------------------

def parse_tasks(html: str, class_id: str = "") -> list[dict]:
    soup = BeautifulSoup(html, "lxml")
    tasks = []

    # Task rows use class "fusion-card-item" — each contains one task link + optional comment
    rows = soup.find_all(class_="fusion-card-item")
    for row in rows:
        # Find the task title link
        a = row.find("a", href=re.compile(r'/core_tasks/\d+'))
        if not a:
            continue
        href = a["href"]
        task_id_match = re.search(r'/core_tasks/(\d+)', href)
        if not task_id_match:
            continue
        task_id = task_id_match.group(1)
        title = a.get_text(strip=True)
        if not title or title in ("Submit Coursework", "View Teacher Feedback"):
            continue

        row_text = row.get_text(" ", strip=True)

        # Date — month + day
        date_match = re.search(
            r'(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s+(\d{1,2})',
            row_text, re.I
        )
        date_str = f"{date_match.group(1)} {date_match.group(2)}" if date_match else ""

        # Due day/time
        due_match = re.search(
            r'(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\s+at\s+\d+:\d+\s*[AP]M',
            row_text, re.I
        )
        due_day_time = due_match.group(0) if due_match else ""

        # Tags — label spans
        # ManageBac uses badge/label/tag elements but many of them contain dates, status
        # words, grade numbers, or long concatenations — filter those out aggressively.
        _JUNK_TAG = re.compile(
            r'^\d+$'                                                             # pure numbers (grade scores)
            r'|^[A-Z][a-z]{2}\d{1,2}$'                                          # dates like "Mar31", "Jun2"
            r'|(Summative|Formative|Pending|Submitted|Complete|Incomplete|Not\s+Submitted|N/A)'
            r'|at\s+\d+:\d+'                                                     # contains a time
            , re.I
        )
        tags = []
        for tag_el in row.find_all(class_=re.compile(r'\btag\b|\bbadge\b|\blabel\b', re.I)):
            t = tag_el.get_text(strip=True)
            if t and len(t) <= 40 and not _JUNK_TAG.search(t):
                tags.append(t)

        task_type = "Summative" if "Summative" in row_text else "Formative" if "Formative" in row_text else ""
        criterion_tags = list(dict.fromkeys(tags))  # deduplicate, preserve order

        # Status — read it from the dedicated status badge (span.badge-label),
        # NOT a substring scan of the whole row. Titles, tags and teacher
        # comments routinely contain words like "submitted" or "complete", and
        # the old scan matched those by accident — e.g. a "Not Submitted" task
        # matched "Submitted", and nothing ever came back "Not Submitted". The
        # badge carries ManageBac's own status text verbatim, with no collisions.
        status = ""
        status_badge = row.find("span", class_="badge-label")
        if status_badge:
            status = status_badge.get_text(" ", strip=True)

        has_submission_box = bool(row.find(string=re.compile(r'Submit Coursework', re.I)))

        # Grades — pattern "A: 7 8" or "B: N/A"
        grades = {}
        for match in re.finditer(r'\b([A-D]):\s*(\d+)\s+(\d+)', row_text):
            criterion, score, max_score = match.groups()
            grades[criterion] = {"score": int(score), "max": int(max_score)}
        for match in re.finditer(r'\b([A-D]):\s*N/A\b', row_text, re.IGNORECASE):
            criterion = match.group(1).upper()
            if criterion not in grades:
                grades[criterion] = {"score": None, "max": None}

        # Teacher comment — inside div.assessment-comments > div.show-more > div.fix-body-margins
        comment = ""
        comment_block = row.find(class_="assessment-comments")
        if comment_block:
            content_div = comment_block.find(class_=re.compile(r'fix-body-margins|fr-element|rte', re.I))
            if content_div:
                comment = _html_to_markdown(content_div)
            else:
                comment = comment_block.get_text(" ", strip=True)

        tasks.append({
            "id": task_id,
            "title": title,
            "url": f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}" if class_id else "",
            "date": date_str,
            "due_day_time": due_day_time,
            "type": task_type,
            "tags": criterion_tags,
            "status": status,
            "has_submission_box": has_submission_box,
            "grades": grades,
            "teacher_comment": comment,
        })

    return tasks


async def fetch_tasks(class_id: str) -> list[dict]:
    cache_key = f"get_tasks:{class_id}"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached

    async with await get_client() as client:
        r = await authed_get(client, f"/student/classes/{class_id}/core_tasks")

    result = parse_tasks(r.text, class_id)
    cache.set(cache_key, result, "get_tasks")
    return result


# ---------------------------------------------------------------------------
# Upcoming / overdue (consolidated deadlines across ALL classes)
# ---------------------------------------------------------------------------

_UPCOMING_STATUSES = [
    "Not Assessed Yet", "Not Submitted", "Submitted",
    "Pending", "Complete", "Incomplete", "Overdue",
]


def parse_upcoming(html: str) -> list[dict]:
    """
    Parse /student/tasks_and_deadlines — the consolidated list of every task
    across all classes, grouped by day. This is the authoritative source for
    'what is due today / this week', far more reliable than per-class crawling.
    """
    soup = BeautifulSoup(html, "lxml")
    container = soup.find("div", class_="js-tasks")
    items: list[dict] = []
    if not container:
        return items

    current_day = ""
    for el in container.children:
        if getattr(el, "name", None) is None:
            continue
        if el.name == "p" and el.get("role") == "heading":
            current_day = el.get_text(" ", strip=True)   # e.g. "Today - Tuesday, Jun 2"
        elif el.name == "section":
            for tile in el.find_all("div", class_="f-task-tile"):
                link = tile.find("a", class_="f-tile__title-link") or \
                       tile.find("a", href=re.compile(r"/core_tasks/\d+"))
                href = link.get("href", "") if link else ""
                m = re.search(r"/classes/(\d+)/core_tasks/(\d+)", href)
                if not m:
                    continue
                cid, tid = m.group(1), m.group(2)
                title = link.get_text(" ", strip=True)

                class_a = tile.find("a", href=re.compile(r"/student/classes/\d+$"))
                class_name = class_a.get_text(" ", strip=True) if class_a else ""

                tile_text = tile.get_text(" ", strip=True)
                due_m = re.search(r"[A-Z][a-z]{2}\s+\d{1,2},\s*\d{1,2}:\d{2}\s*[AP]M", tile_text)
                due = due_m.group(0) if due_m else ""

                task_type = ("Summative" if "Summative" in tile_text
                             else "Formative" if "Formative" in tile_text else "")
                status = next((s for s in _UPCOMING_STATUSES if s in tile_text), "")
                needs_submission = "Submit Coursework" in tile_text

                items.append({
                    "title": title,
                    "url": f"{_base()}/student/classes/{cid}/core_tasks/{tid}",
                    "class_name": class_name,
                    "class_id": cid,
                    "due": due,
                    "due_group": current_day,
                    "type": task_type,
                    "status": status,
                    "needs_submission": needs_submission,
                })
    return items


async def fetch_upcoming(view: str = "upcoming") -> dict:
    view = view if view in ("upcoming", "overdue", "past") else "upcoming"
    cache_key = f"get_upcoming:{view}"
    items = cache.get(cache_key)
    if items is None:
        async with await get_client() as client:
            r = await authed_get(client, f"/student/tasks_and_deadlines?view={view}")
        items = parse_upcoming(r.text)
        cache.set(cache_key, items, "get_tasks")   # reuse the 10-min tasks TTL
    return {"current": _now_info(), "view": view, "tasks": items}


async def tag_search(tag: str = "", class_id: str = "") -> dict:
    """
    Find tasks matching a tag/type across all classes (or one class).
    Matches `tag` (case-insensitive) against each task's type (Summative/
    Formative) and its tags (Criterion A, Homework, Test, ...). Reuses the
    cached per-class task lists, so it's fast.
    """
    import asyncio as _asyncio
    classes = await fetch_classes()
    name_by_id = {c["id"]: c["name"] for c in classes}
    ids = [class_id] if class_id else [c["id"] for c in classes]

    task_lists = await _asyncio.gather(
        *[fetch_tasks(cid) for cid in ids], return_exceptions=True
    )

    tagl = tag.lower().strip()
    matches: list[dict] = []
    for cid, tasks in zip(ids, task_lists):
        if isinstance(tasks, Exception):
            continue
        for t in tasks:
            haystack = [t.get("type", "")] + list(t.get("tags", []))
            if not tagl or any(tagl in h.lower() for h in haystack):
                matches.append({
                    "title": t.get("title", ""),
                    "url": t.get("url", ""),
                    "class_name": name_by_id.get(cid, ""),
                    "class_id": cid,
                    "type": t.get("type", ""),
                    "tags": t.get("tags", []),
                    "status": t.get("status", ""),
                    "date": t.get("date", ""),
                    "due_day_time": t.get("due_day_time", ""),
                    "grades": t.get("grades", {}),
                })

    # Cap the number of returned tasks so a broad tag (e.g. "Formative" across
    # all classes) can't flood the context. `count` still reports the true total.
    _TAG_SEARCH_CAP = 40
    total = len(matches)
    out = {
        "query": tag,
        "scope": name_by_id.get(class_id, class_id) if class_id else "all classes",
        "count": total,
        "tasks": matches[:_TAG_SEARCH_CAP],
    }
    if total > _TAG_SEARCH_CAP:
        out["note"] = (f"Showing the first {_TAG_SEARCH_CAP} of {total} matches. "
                       f"Narrow the search with a class_id or a more specific tag.")
    return out


async def fetch_grades(class_id: str = "") -> dict:
    """
    Consolidated grades across all classes (or one). Built from the per-task
    criterion scores already scraped by get_tasks, so it's fast.
    For each class: a per-criterion summary (latest, best, average, out_of) plus
    the list of graded tasks with their scores and teacher comments.
    """
    import asyncio as _asyncio
    classes = await fetch_classes()
    name_by_id = {c["id"]: c["name"] for c in classes}
    ids = [class_id] if class_id else [c["id"] for c in classes]

    task_lists = await _asyncio.gather(
        *[fetch_tasks(cid) for cid in ids], return_exceptions=True
    )

    out_classes = []
    fetch_errors = []
    for cid, tasks in zip(ids, task_lists):
        if isinstance(tasks, Exception):
            # Surface the failure instead of silently dropping the class — a
            # missing class here is otherwise indistinguishable from "no grades".
            fetch_errors.append({"class_id": cid, "class_name": name_by_id.get(cid, ""),
                                 "error": str(tasks)})
            continue
        graded = [t for t in tasks if t.get("grades")]   # tasks are newest-first
        if not graded:
            continue

        crit: dict[str, dict] = {}
        for t in graded:
            for k, v in t["grades"].items():
                score = v.get("score")
                if score is None:
                    continue
                entry = crit.setdefault(k, {"scores": [], "out_of": v.get("max")})
                entry["scores"].append(score)
                if v.get("max"):
                    entry["out_of"] = v["max"]

        criteria = {}
        for k in sorted(crit):
            s = crit[k]["scores"]
            criteria[k] = {
                "latest": s[0] if s else None,        # newest graded task first
                "best": max(s) if s else None,
                "average": round(sum(s) / len(s), 1) if s else None,
                "out_of": crit[k]["out_of"],
                "count": len(s),
            }

        entry = {
            "class_name": name_by_id.get(cid, ""),
            "class_id": cid,
            "criteria": criteria,
        }
        # The per-criterion summary above is all the all-classes view needs (and
        # all predicted-grades needs). Only attach the full graded-task list —
        # with teacher comments — when scoped to ONE class. Shipping every task
        # for all ~13 graded classes is what made the payload so large it got
        # truncated by the connector and came back looking empty.
        if class_id:
            entry["graded_tasks"] = [
                {"title": t.get("title", ""), "url": t.get("url", ""),
                 "type": t.get("type", ""), "date": t.get("date", ""),
                 "grades": t["grades"], "teacher_comment": t.get("teacher_comment", "")}
                for t in graded
            ]
        out_classes.append(entry)

    result = {
        "scope": name_by_id.get(class_id, class_id) if class_id else "all classes",
        "classes": out_classes,
    }
    if fetch_errors:
        result["fetch_errors"] = fetch_errors
    return result


async def prewarm(user) -> None:
    """
    Warm a freshly-enrolled user's cache in the background so their first
    question is instant. Fetches the things people ask for first — classes,
    timetable, the consolidated upcoming list, and every class's task list —
    concurrently. Best-effort: any failure is swallowed (the data just loads
    lazily on first use instead).
    """
    import asyncio as _asyncio
    from .context import set_current_user
    set_current_user(user)   # this task gets its own isolated context
    try:
        classes = await fetch_classes()
        class_ids = [c["id"] for c in classes]
        await _asyncio.gather(
            fetch_timetable(),
            fetch_upcoming("upcoming"),
            *[fetch_tasks(cid) for cid in class_ids],
            return_exceptions=True,
        )
        print(f"[prewarm] cached {len(class_ids)} classes for {user.label}", flush=True)
    except Exception as e:
        print(f"[prewarm] failed for {user.label}: {e}", flush=True)


# ---------------------------------------------------------------------------
# Task detail
# ---------------------------------------------------------------------------

def parse_task_detail(html: str, class_id: str, task_id: str) -> dict:
    soup = BeautifulSoup(html, "lxml")

    # Task title — typically in h2.page-title or the first prominent heading
    title = ""
    for sel in ["h2.page-title", "h1.page-title", ".page-title", "h1", "h2"]:
        el = soup.select_one(sel)
        if el:
            t = el.get_text(strip=True)
            # Skip generic headings like "Description", "Resources"
            if t and t.lower() not in ("description", "resources", "discussions", "task history"):
                title = t
                break

    # Description block — find <h4>Description</h4> then get next sibling div.show-more
    # Inside that: div.fix-body-margins (the actual content)
    desc_heading = soup.find(["h4", "h3", "h2"], string=re.compile(r'^Description$', re.I))
    desc_block = None
    if desc_heading:
        sib = desc_heading.find_next_sibling()
        if sib:
            # May be div.show-more or directly the content div
            inner = sib.find(class_=re.compile(r'fix-body-margins|fr-element', re.I))
            desc_block = inner if inner else sib

    desc_text = ""
    desc_links = []
    desc_files = []

    if desc_block:
        desc_text = _html_to_markdown(desc_block)

        for a in desc_block.find_all("a", href=True):
            href = a.get("href", "")
            a_classes = a.get("class", [])

            # Embedded file — ManageBac uses <a class="fr-file" data-name="...">
            if "fr-file" in a_classes or a.get("data-name"):
                name = a.get("data-name") or a.get_text(strip=True)
                size_el = a.find(class_="fr-file-size")
                size = size_el.get_text(strip=True) if size_el else ""
                if name:
                    desc_files.append({"name": name, "size": size, "url": href})

            # External link (Google Docs, YouTube, etc.)
            elif href.startswith("http") and "managebac.com" not in href:
                text = a.get_text(strip=True)
                # Use the href as text if the anchor text is just the URL itself
                desc_links.append({"text": text if text != href else href, "url": href})

        # Also catch bare URLs typed as plain text (not wrapped in <a>)
        _URL_RE = re.compile(r'https?://[^\s<>"\']+')
        linked_urls = {d["url"] for d in desc_links}
        for text_node in desc_block.find_all(string=_URL_RE):
            # Skip if inside an <a> tag (already captured above)
            if text_node.parent and text_node.parent.name == "a":
                continue
            for url in _URL_RE.findall(str(text_node)):
                if url not in linked_urls and "managebac.com" not in url:
                    desc_links.append({"text": url, "url": url})
                    linked_urls.add(url)

    # Resources section (teacher-posted files on the task)
    resources = []
    resources_section = soup.find(string=re.compile(r'^Resources$', re.I))
    if resources_section:
        resources_container = resources_section.find_parent()
        while resources_container and not resources_container.find_all(class_=re.compile(r'file|attachment', re.I)):
            resources_container = resources_container.find_parent()
        if resources_container:
            for post in resources_container.find_all(class_=re.compile(r'post|entry|item', re.I)):
                teacher_el = post.find(class_=re.compile(r'author|name|user', re.I))
                teacher = teacher_el.get_text(strip=True) if teacher_el else ""
                label_el = post.find(["p", "strong", "b"])
                label = label_el.get_text(strip=True) if label_el else ""
                files = []
                for f in post.find_all(class_=re.compile(r'file|attachment', re.I)):
                    fname = f.find("a")
                    fsize = f.find(string=re.compile(r'\d+\s*(KB|MB)', re.I))
                    files.append({
                        "name": fname.get_text(strip=True) if fname else "",
                        "size": str(fsize).strip() if fsize else "",
                    })
                if teacher or files:
                    resources.append({"teacher": teacher, "label": label, "files": files})

    # Dropbox / submitted files — each is a <tr class="file"> inside the dropbox table
    submitted_files = []
    for tr in soup.find_all("tr", class_="file"):
        # Filename — in <a class="text-break"> or <div class="details"> > <a>
        file_link = tr.find("a", class_="text-break") or tr.find(class_="details", recursive=True)
        if file_link and file_link.name != "a":
            file_link = file_link.find("a")
        fname = file_link.get_text(strip=True) if file_link else ""
        if not fname:
            continue

        # The real, downloadable file URL — a signed CDN link in the filename
        # anchor's href. THIS is what get_file_content needs. (The old code kept
        # only the preview-modal token below, which is an HTML popup, not a file,
        # so reading a submitted file 404'd.)
        download_url = ""
        if file_link and file_link.name == "a":
            download_url = file_link.get("href", "")
        if not download_url:
            asset = tr.find("a", href=re.compile(r"/uploads/asset/|cdn\.[^/]*managebac", re.I))
            download_url = asset.get("href", "") if asset else ""

        # Upload timestamp — in a <label> inside the row
        uploaded_at = ""
        label = tr.find("label")
        if label:
            t = label.get_text(strip=True)
            up_match = re.search(r'Uploaded (.+)', t)
            if up_match:
                uploaded_at = up_match.group(1)

        # Teacher feedback token (the preview-modal popup — not a file download)
        feedback_btn = tr.find("a", attrs={"data-pdf-preview-url-value": True})
        feedback_token = feedback_btn["data-pdf-preview-url-value"] if feedback_btn else None

        submitted_files.append({
            "name": fname,
            "url": download_url,          # pass this to get_file_content to read it
            "uploaded_at": uploaded_at,
            "teacher_feedback_token": feedback_token,
        })

    # Task history from sidebar
    history = {}
    history_section = soup.find(string=re.compile(r'Task History', re.I))
    if history_section:
        history_container = history_section.find_parent()
        text = history_container.get_text(" ", strip=True) if history_container else ""
        created = re.search(r'Created (.+?)(?:Reminder|Last Updated|$)', text)
        reminder = re.search(r'Reminder Sent (.+?)(?:Last Updated|$)', text)
        updated = re.search(r'Last Updated (.+?)$', text)
        if created:
            history["created_at"] = created.group(1).strip()
        if reminder:
            history["reminder_sent_at"] = reminder.group(1).strip()
        if updated:
            history["last_updated_at"] = updated.group(1).strip()

    return {
        "class_id": class_id,
        "task_id": task_id,
        "url": f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}",
        **({"title": title} if title else {}),
        "description": {
            "text": desc_text,
            "links": desc_links,
            "embedded_files": desc_files,
        },
        "resources": resources,
        "submitted_files": submitted_files,
        "task_history": history,
    }


def parse_discussions(html: str, class_id: str, task_id: str) -> list[dict]:
    """
    Parse discussion posts from /core_tasks/{id}/discussions.

    Each post is a div.discussion containing:
      div.author          → poster's name
      div.date.gray-text  → "Posted on Friday, Nov 14, 2025 at 4:00 PM"
      div.body.pt-3
        div.fix-body-margins  → message text + any embedded links/files

    Replies (if any) appear as div.reply elements beneath the main post.
    """
    soup = BeautifulSoup(html, "lxml")
    posts = []

    for post in soup.find_all("div", class_="discussion"):
        post_id_attr = post.get("id", "")  # e.g. "discussion_27539214"
        post_id = post_id_attr.replace("discussion_", "")

        # Author
        author_el = post.find("div", class_="author")
        author = author_el.get_text(strip=True) if author_el else ""

        # Date — "Posted on Friday, Nov 14, 2025 at 4:00 PM"
        date_el = post.find("div", class_="date")
        date_raw = date_el.get_text(strip=True) if date_el else ""
        date = re.sub(r'^Posted on\s*', '', date_raw).strip()

        # Body — text + links
        # Use the whole div.body container so text nodes outside fix-body-margins are included
        body_div = post.find("div", class_="body")
        body_text = ""
        body_links = []
        if body_div:
            body_text = _html_to_markdown(body_div)
            for a in body_div.find_all("a", href=True):
                href = a.get("href", "")
                text = a.get_text(strip=True)
                if href.startswith("http"):
                    body_links.append({"text": text or href, "url": href})

        # Replies — div.reply elements anywhere in the post
        replies = []
        for reply in post.find_all("div", class_="reply"):
            if "form" in reply.get("class", []):
                continue  # skip the reply input form
            r_author_el = reply.find(class_=re.compile(r'author|name', re.I))
            r_author = r_author_el.get_text(strip=True) if r_author_el else ""
            r_date_el = reply.find("div", class_="date")
            r_date_raw = r_date_el.get_text(strip=True) if r_date_el else ""
            r_date = re.sub(r'^Posted on\s*', '', r_date_raw).strip()
            r_body_div = reply.find("div", class_="body")
            r_text = ""
            r_links = []
            if r_body_div:
                r_text = _html_to_markdown(r_body_div)
                for a in r_body_div.find_all("a", href=True):
                    href = a.get("href", "")
                    t = a.get_text(strip=True)
                    if href.startswith("http"):
                        r_links.append({"text": t or href, "url": href})
            if r_author or r_text:
                replies.append({
                    "author": r_author,
                    "posted_at": r_date,
                    "text": r_text,
                    "links": r_links,
                })

        if author or body_text:
            posts.append({
                "id": post_id,
                "author": author,
                "posted_at": date,
                "text": body_text,
                "links": body_links,
                "replies": replies,
            })

    return posts


def parse_theme(html: str) -> str:
    """Return the student's selected ManageBac colour theme.

    ManageBac stamps the chosen theme onto the <body> class of every page as
    `theme-<name>` (blue / orange / red / plum / teal). Defaults to 'teal'.
    """
    m = re.search(r"theme-(blue|orange|red|plum|teal)\b", html)
    return m.group(1) if m else "teal"


async def fetch_task_detail(class_id: str, task_id: str) -> dict:
    cache_key = f"get_task_detail:{class_id}:{task_id}"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached

    async with await get_client() as client:
        r = await authed_get(client, f"/student/classes/{class_id}/core_tasks/{task_id}")
        r_disc = await authed_get(client, f"/student/classes/{class_id}/core_tasks/{task_id}/discussions")

    result = parse_task_detail(r.text, class_id, task_id)
    result["discussions"] = parse_discussions(r_disc.text, class_id, task_id)
    result["theme"] = parse_theme(r.text)  # free: read from the page we already fetched
    cache.set(cache_key, result, "get_task_detail")
    return result


# ---------------------------------------------------------------------------
# Files
# ---------------------------------------------------------------------------

def parse_files(html: str) -> list[dict]:
    import json as _json
    # lxml strips data-ec3-info attributes — use html.parser here
    soup = BeautifulSoup(html, "html.parser")
    files = []

    # Class files use <div class="row file" data-ec3-info='{...}'> elements
    # data-ec3-info JSON has: name, file_size, created_at, updated_at
    for row in soup.find_all(attrs={"data-ec3-info": True}):
        try:
            info = _json.loads(row["data-ec3-info"])
        except Exception:
            continue
        name = info.get("name", "")
        if not name:
            continue

        file_size_bytes = info.get("file_size", 0)
        if file_size_bytes:
            if file_size_bytes >= 1_048_576:
                size = f"{file_size_bytes / 1_048_576:.1f} MB"
            else:
                size = f"{file_size_bytes / 1024:.1f} KB"
        else:
            size = ""

        # Author — inside the row's author element
        author_el = row.find(class_=re.compile(r'author|uploader|user', re.I))
        if not author_el:
            author_el = row.find("a", href=re.compile(r'/users/', re.I))
        author = author_el.get_text(strip=True) if author_el else ""

        # Date and download URL from JSON
        date = info.get("updated_at", info.get("created_at", ""))
        url = info.get("download_url", "")

        files.append({
            "name": name,
            "size": size,
            "url": url,
            "uploaded_by": author,
            "uploaded_at": date,
        })

    # Fallback: <tr class="file"> rows (used in some class file views)
    if not files:
        for tr in soup.find_all("tr", class_="file"):
            name_link = tr.find("a", class_="text-break")
            if not name_link:
                continue
            name = name_link.get_text(strip=True)
            row_text = tr.get_text(" ", strip=True)
            size_match = re.search(r'(\d+(?:\.\d+)?\s*(?:KB|MB|GB))', row_text, re.I)
            size = size_match.group(1) if size_match else ""
            author_match = re.search(r'by\s+(.+?)(?:\s+\d|$)', row_text)
            author = author_match.group(1).strip() if author_match else ""
            files.append({"name": name, "size": size, "uploaded_by": author, "uploaded_at": ""})

    return files


async def fetch_files(class_id: str) -> list[dict]:
    cache_key = f"get_files:{class_id}"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached

    async with await get_client() as client:
        r = await authed_get(client, f"/student/classes/{class_id}/files")

    result = parse_files(r.text)
    cache.set(cache_key, result, "get_files")
    return result


# ---------------------------------------------------------------------------
# Journal
# ---------------------------------------------------------------------------

def parse_journal(html: str) -> list[dict]:
    soup = BeautifulSoup(html, "lxml")
    entries = []

    # Each journal entry is div.journal-evidence (e.g. id="evidence-58103436")
    # Structure:
    #   h4.title
    #     span.padding-right  → date "March 20, 2026"
    #     small               → time "1:33 AM"
    #     small.labels-set    → learning outcomes
    #     div.actions         → star / edit / delete links
    #       a[data-star]      → star toggle; href contains /star; if starred the SVG fill changes
    #   div.body
    #     div.fix-body-margins  → reflection text + any embedded links/files
    for entry in soup.find_all("div", class_="journal-evidence"):
        entry_id = entry.get("id", "")  # e.g. "evidence-58103436"

        # Date
        date_span = entry.find("span", class_="padding-right")
        date = date_span.get_text(strip=True) if date_span else ""
        if not date:
            continue

        # Time (optional)
        time_el = entry.find("small", class_=False)
        time_str = time_el.get_text(strip=True) if time_el else ""

        # Learning outcomes / labels — skip meta-labels like "Read-only"
        outcomes = []
        for label in entry.find_all("div", class_="label-outcome"):
            t = label.get_text(strip=True)
            if t and not re.match(r'^read.?only$', t, re.I):
                outcomes.append(t)

        # Starred — the star <a> has data-star attribute; if the path contains "/unstar"
        # or the SVG has a filled colour it's starred. Simplest signal: look for data-unstar
        # text = "Unstar" (meaning it IS currently starred, click to unstar).
        star_a = entry.find("a", attrs={"data-star": True})
        is_starred = False
        if star_a:
            # If button currently says "Unstar" → already starred
            is_starred = star_a.get("data-unstar", "").lower() == "unstar" and \
                         star_a.get("data-star", "").lower() == "star"
            # Fallback: check href — starred entries have /star in href still (toggle),
            # but we can also check the SVG fill colour
            svg = star_a.find("svg")
            if svg:
                path = svg.find("path")
                if path:
                    fill = path.get("fill", "")
                    # Filled gold/yellow star = starred; outline grey = not starred
                    is_starred = fill not in ("", "#859bbb", "none")

        is_read_only = bool(entry.find(string=re.compile(r'Read.?only', re.I)))

        # Body text and links
        body_div = entry.find("div", class_="fix-body-margins")
        body_text = _html_to_markdown(body_div) if body_div else ""

        links = []
        files = []
        if body_div:
            for a in body_div.find_all("a", href=True):
                href = a.get("href", "")
                text = a.get_text(strip=True)
                a_classes = a.get("class", [])
                if "fr-file" in a_classes or a.get("data-name"):
                    # Embedded file attachment
                    name = a.get("data-name") or text
                    size_el = a.find(class_="fr-file-size")
                    size = size_el.get_text(strip=True) if size_el else ""
                    files.append({"name": name, "size": size})
                elif href.startswith("http"):
                    links.append({"text": text or href, "url": href})

        entries.append({
            "id": entry_id.replace("evidence-", ""),
            "date": date,
            "time": time_str,
            "learning_outcomes": outcomes,
            "is_starred": is_starred,
            "is_read_only": is_read_only,
            "body": body_text,
            "links": links,
            "files": files,
        })

    return entries


async def fetch_journal(class_id: str) -> list[dict]:
    cache_key = f"get_journal:{class_id}"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached

    async with await get_client() as client:
        r = await authed_get(client, f"/student/classes/{class_id}/learner_portfolio/reflections")

    # If no journal for this class — return empty
    if r.status_code == 404 or "/login" in str(r.url):
        return []
    if "learner_portfolio" not in str(r.url) and "Journal" not in r.text:
        return []

    result = parse_journal(r.text)
    cache.set(cache_key, result, "get_journal")
    return result


# ---------------------------------------------------------------------------
# Units
# ---------------------------------------------------------------------------

def parse_units(html: str, class_id: str) -> list[dict]:
    """
    Parse the /units page to get each unit's id, title, start date,
    duration, and current status.  IB detail fields are filled in later
    by parse_unit_popup().
    """
    soup = BeautifulSoup(html, "lxml")
    units = []

    prefix = f"ib_class_{class_id}_core_unit_"
    for div in soup.find_all("div", id=re.compile(rf"^ib_class_{class_id}_core_unit_\d+$")):
        div_id = div.get("id", "")
        unit_id = div_id.replace(prefix, "")
        if not unit_id.isdigit():
            continue

        # Skip weekly-planner duplicates
        weekly_id = f"ib_class_{class_id}_weekly_core_unit_{unit_id}"
        if soup.find("div", id=weekly_id) and div_id != weekly_id:
            pass  # keep the canonical one

        unit_comp = div.find("div", class_="unit-component")
        if not unit_comp:
            continue

        # Title
        h4 = unit_comp.find("p", class_="h4")
        title_a = h4.find("a") if h4 else None
        title = title_a.get_text(strip=True) if title_a else ""
        if not title:
            continue

        # Start date — "Starts\n W4 Nov" → "W4 Nov"
        start_el = unit_comp.find("span", class_="label-start")
        start = ""
        if start_el:
            inner_span = start_el.find("span")
            if inner_span:
                raw = inner_span.get_text(" ", strip=True)
                start = re.sub(r'^Starts\s*', '', raw).strip()

        # Duration — "4 Weeks" (after stripping the clock SVG)
        duration_el = unit_comp.find("div", class_="unit-duration")
        duration = ""
        if duration_el:
            dur_clone = BeautifulSoup(str(duration_el), "lxml").find("div")
            for svg in dur_clone.find_all("svg"):
                svg.decompose()
            duration = dur_clone.get_text(" ", strip=True)

        # Status: current / completed / upcoming
        comp_classes = unit_comp.get("class", [])
        if "current" in comp_classes:
            status = "current"
        elif "completed" in comp_classes:
            status = "completed"
        else:
            status = "upcoming"

        units.append({
            "id": unit_id,
            "title": title,
            "start": start,
            "duration": duration,
            "status": status,
            "url": f"{_base()}/student/classes/{class_id}/units/{unit_id}/presentations",
        })

    return units


def parse_unit_popup(html: str) -> dict:
    """
    Parse the unit popup fragment (/units/{id}/popup) to extract all IB
    framework fields: global_context, key_concepts, related_concepts,
    conceptual_understanding, statement_of_inquiry, inquiry_questions, atl_skills.
    Also extracts start_date and duration from the <dl> block.
    """
    from bs4 import NavigableString
    soup = BeautifulSoup(html, "lxml")
    result: dict[str, Any] = {}

    # Basic metadata from dt/dd pairs
    for dt in soup.find_all("dt"):
        dd = dt.find_next_sibling("dd")
        label = dt.get_text(" ", strip=True)
        value = dd.get_text(" ", strip=True) if dd else ""
        if "Start date" in label:
            result["start_date"] = value
        elif "Duration" in label:
            result["duration"] = re.sub(r'\s+', ' ', value).strip()

    # IB framework sections
    for section in soup.find_all("div", class_="section"):
        header = section.find("div", class_="unit-component-header")
        if not header:
            continue
        header_text = header.get_text(strip=True)

        if "Global Context" in header_text:
            contexts = []
            for flex in section.find_all("div", class_="flex"):
                clone = BeautifulSoup(str(flex), "lxml").find("div")
                for img in clone.find_all("img"):
                    img.decompose()
                t = clone.get_text(" ", strip=True)
                if t:
                    contexts.append(t)
            result["global_context"] = contexts

        elif "Key Concept" in header_text:
            concepts = []
            for p in section.find_all("p"):
                clone = BeautifulSoup(str(p), "lxml").find("p")
                if not clone:
                    continue
                for img in clone.find_all("img"):
                    img.decompose()
                name_el = clone.find("strong")
                name = name_el.get_text(strip=True) if name_el else ""
                full = clone.get_text(" ", strip=True)
                if name and name in full:
                    definition = full[full.index(name) + len(name):].strip()
                else:
                    definition = full
                    name = full
                if name:
                    concepts.append({"name": name, "definition": definition})
            result["key_concepts"] = concepts

        elif "Related Concept" in header_text:
            for p in section.find_all("p"):
                t = p.get_text(" ", strip=True)
                if t:
                    result["related_concepts"] = [c.strip() for c in t.split(",") if c.strip()]
                    break

        elif "Conceptual Understanding" in header_text:
            body = section.find(class_="fix-body-margins")
            if body:
                result["conceptual_understanding"] = _html_to_markdown(body)

        elif "Statement of Inquiry" in header_text:
            body = section.find(class_="fix-body-margins")
            if body:
                result["statement_of_inquiry"] = _html_to_markdown(body)

        elif "Inquiry Question" in header_text:
            questions = []
            for p in section.find_all("p", class_="mb-1"):
                label_el = p.find("span", class_=re.compile(r"\blabel\b"))
                q_type = label_el.get_text(strip=True) if label_el else ""
                # Get only the direct text nodes (not inside the span)
                q_text = "".join(
                    str(s) for s in p.children
                    if isinstance(s, NavigableString)
                ).strip()
                if q_type or q_text:
                    questions.append({"type": q_type, "question": q_text})
            result["inquiry_questions"] = questions

        elif "ATL" in header_text:
            skills = []
            for flex in section.find_all("div", class_="flex"):
                clone = BeautifulSoup(str(flex), "lxml").find("div")
                for img in clone.find_all("img"):
                    img.decompose()
                t = clone.get_text(" ", strip=True)
                if t:
                    skills.append(t)
            result["atl_skills"] = skills

    return result


async def fetch_units(class_id: str) -> list[dict]:
    cache_key = f"get_units:{class_id}"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached

    import asyncio as _asyncio

    async with await get_client() as client:
        r = await authed_get(client, f"/student/classes/{class_id}/units")
        units = parse_units(r.text, class_id)

        if not units:
            return []

        # Fetch all unit popups concurrently in the same session
        popup_responses = await _asyncio.gather(*[
            authed_get(client, f"/student/classes/{class_id}/units/{u['id']}/popup")
            for u in units
        ])

    for unit, popup_r in zip(units, popup_responses):
        ib = parse_unit_popup(popup_r.text)
        unit.update(ib)

    cache.set(cache_key, units, "get_units")
    return units


# ---------------------------------------------------------------------------
# File download (raw bytes, no conversion)
# ---------------------------------------------------------------------------

_MAX_FILE_BYTES = 20 * 1024 * 1024   # 20 MB hard limit (Anthropic API PDF cap)
_FILES_CACHE_DIR = None               # set lazily from config


def _file_cache_paths(url: str):
    """Return (meta_path, data_path) for the given URL."""
    import hashlib
    from .config import DATA_DIR
    cache_dir = DATA_DIR / "files"
    cache_dir.mkdir(exist_ok=True)
    h = hashlib.md5(url.encode()).hexdigest()
    return cache_dir / f"{h}.json", cache_dir / f"{h}.bin"


async def fetch_file_bytes(url: str) -> dict[str, Any]:
    """
    Download a ManageBac attachment using the authenticated session and
    return the raw bytes + metadata.  Results are cached to disk for 1 hour.

    Returns:
        {content_type, size_bytes, data (bytes), error (None on success)}
    """
    import json as _json, time as _time

    meta_path, data_path = _file_cache_paths(url)

    # --- disk cache hit ---
    if meta_path.exists() and data_path.exists():
        try:
            meta = _json.loads(meta_path.read_text())
            if _time.time() < meta.get("expires_at", 0):
                return {
                    "content_type": meta["content_type"],
                    "size_bytes": meta["size_bytes"],
                    "data": data_path.read_bytes(),
                    "error": None,
                }
        except Exception:
            pass  # stale / corrupt cache — fall through to re-download

    async with await get_client() as client:
        try:
            r = await authed_get(client, url)
        except Exception as e:
            return {"content_type": "", "size_bytes": 0, "data": b"", "error": f"Download failed: {e}"}

    if r.status_code != 200:
        return {"content_type": "", "size_bytes": 0, "data": b"", "error": f"HTTP {r.status_code}"}

    content_type = r.headers.get("content-type", "application/octet-stream").split(";")[0].strip()
    data = r.content

    if len(data) > _MAX_FILE_BYTES:
        return {
            "content_type": content_type,
            "size_bytes": len(data),
            "data": b"",
            "error": f"File is {len(data) / 1_048_576:.1f} MB — exceeds the 20 MB limit.",
        }

    # --- write disk cache ---
    try:
        meta_path.write_text(_json.dumps({
            "content_type": content_type,
            "size_bytes": len(data),
            "expires_at": int(_time.time()) + 3600,
        }))
        data_path.write_bytes(data)
    except Exception:
        pass  # cache write failure is non-fatal

    return {"content_type": content_type, "size_bytes": len(data), "data": data, "error": None}


_MAX_EXTRACT_CHARS = 40_000   # cap so a huge PDF can't blow up the AI's context


async def fetch_file_readable(url: str) -> dict[str, Any]:
    """
    Download an attachment and return something an AI can actually use without
    blowing up its context: extracted TEXT for documents (PDF/DOCX/text), or an
    image for image files. Avoids dumping multi-MB base64 blobs into the chat.

    Returns one of:
      {"kind": "text",  "text": "...", "truncated": bool, "content_type": ...}
      {"kind": "image", "data_b64": "...", "content_type": "image/..."}
      {"kind": "error", "error": "..."}
    """
    import io, base64 as _b64
    f = await fetch_file_bytes(url)
    if f["error"]:
        return {"kind": "error", "error": f["error"]}

    data = f["data"]
    ct = (f["content_type"] or "").lower()
    ext = url.split("?")[0].rsplit(".", 1)[-1].lower() if "." in url.split("?")[0] else ""

    # Images — return as image content (AI can view it; usually small enough)
    if ct.startswith("image/") or ext in ("png", "jpg", "jpeg", "gif", "webp"):
        return {"kind": "image", "data_b64": _b64.standard_b64encode(data).decode(),
                "content_type": ct or f"image/{ext}"}

    text = ""
    try:
        if "pdf" in ct or ext == "pdf":
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text()
                    if t and t.strip():
                        parts.append(f"[Page {i}]\n{t.strip()}")
            text = "\n\n".join(parts)
        elif "word" in ct or "officedocument" in ct or ext in ("docx",):
            import docx as _docx
            doc = _docx.Document(io.BytesIO(data))
            chunks = [p.text for p in doc.paragraphs if p.text.strip()]
            # include table cells too (templates are often tables)
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        chunks.append(" | ".join(cells))
            text = "\n".join(chunks)
        elif "text" in ct or ext in ("txt", "md", "csv"):
            text = data.decode("utf-8", errors="replace")
        else:
            return {"kind": "error",
                    "error": f"Can't read this file type ({ct or ext or 'unknown'}). "
                             f"Supported: PDF, Word (.docx), text, images."}
    except Exception as e:
        return {"kind": "error", "error": f"Could not extract text: {e}"}

    truncated = len(text) > _MAX_EXTRACT_CHARS
    if truncated:
        text = text[:_MAX_EXTRACT_CHARS] + "\n\n[…truncated]"
    return {"kind": "text", "text": text or "(no readable text found in this file)",
            "truncated": truncated, "content_type": ct}


# ---------------------------------------------------------------------------
# Task file submission
# ---------------------------------------------------------------------------

async def submit_task_file(
    class_id: str,
    task_id: str,
    file_path: str | None = None,
    file_base64: str | None = None,
    filename: str | None = None,
    dry_run: bool = False,
) -> dict[str, Any]:
    """
    Upload a file to a task's dropbox on ManageBac.

    Two ways to provide the file:
      • file_path  — absolute path to a local file (used by Claude Desktop / CLI,
                     where the server runs on the same machine as the file)
      • file_base64 + filename — the file's raw bytes, base64-encoded (used by
                     remote clients like ChatGPT, where there is no shared filesystem)

    Workflow:
      1. GET the dropbox page to get a fresh CSRF token
      2. POST multipart/form-data to the upload endpoint

    Rails endpoint:
      POST /student/classes/{class_id}/core_tasks/{task_id}/dropbox/upload
      _method=patch, X-CSRF-Token header, dropbox[assets_attributes][0][file]

    If dry_run=True, validates everything but does NOT submit — returns a preview.
    """
    from pathlib import Path as _Path
    from bs4 import BeautifulSoup as _BS
    import mimetypes as _mimetypes
    import os as _os
    import base64 as _base64

    # --- resolve the file from either source ---
    if file_base64:
        if not filename:
            return {"success": False, "error": "filename is required when sending file_base64."}
        try:
            file_bytes = _base64.b64decode(file_base64)
        except Exception as e:
            return {"success": False, "error": f"Could not decode file_base64: {e}"}
        upload_name = filename
        source = "(in-memory bytes)"
    elif file_path:
        path = _Path(file_path).expanduser()
        if not path.is_absolute():
            resolved = _Path.home() / path
            if resolved.exists():
                path = resolved
            else:
                return {
                    "success": False,
                    "error": (
                        f"Relative path '{file_path}' not found. "
                        f"Use an absolute path like '/tmp/{path.name}'. "
                        f"Current working directory is: {_os.getcwd()}"
                    ),
                }
        if not path.exists():
            return {
                "success": False,
                "error": (
                    f"File not found: {path}. "
                    "Save the file to an absolute path like /tmp/blank.pdf before calling this tool."
                ),
            }
        if not path.is_file():
            return {"success": False, "error": f"Not a file: {path}"}
        file_bytes = path.read_bytes()
        upload_name = path.name
        source = str(path)
    else:
        return {"success": False, "error": "Provide either file_path or file_base64 + filename."}

    file_size = len(file_bytes)
    if file_size > 500 * 1024 * 1024:
        return {"success": False, "error": f"File too large ({file_size / 1_048_576:.1f} MB). ManageBac limit is 500 MB."}

    mime_type = _mimetypes.guess_type(upload_name)[0] or "application/octet-stream"

    async with await get_client() as client:
        # Fetch dropbox page to get CSRF token and confirm dropbox exists
        dropbox_url = f"/student/classes/{class_id}/core_tasks/{task_id}/dropbox"
        r = await authed_get(client, dropbox_url)

        if r.status_code != 200 or "/login" in str(r.url):
            return {"success": False, "error": "Could not access the dropbox page — check class_id and task_id."}

        soup = _BS(r.text, "lxml")
        csrf_el = soup.find("meta", {"name": "csrf-token"})
        if not csrf_el:
            return {"success": False, "error": "Could not find CSRF token on the dropbox page."}
        csrf_token = csrf_el["content"]

        # Confirm the upload form is present
        form = soup.find("form", id=lambda i: i and "dropbox" in i)
        if not form:
            return {"success": False, "error": "No upload form found — this task may not have a submission dropbox."}

        if dry_run:
            return {
                "dry_run": True,
                "would_submit": {
                    "file": source,
                    "filename": upload_name,
                    "size_bytes": file_size,
                    "mime_type": mime_type,
                    "to_task": f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}",
                    "endpoint": f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}/dropbox/upload",
                },
                "note": "Set dry_run=false to actually submit.",
            }

        upload_url = f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}/dropbox/upload"
        response = await client.post(
            upload_url,
            data={
                "_method": "patch",
                "dropbox[assets_attributes][0][file_cache]": "",
                "commit": "Upload Files",
            },
            files={
                "dropbox[assets_attributes][0][file]": (upload_name, file_bytes, mime_type),
            },
            headers={
                "X-CSRF-Token": csrf_token,
                "X-Requested-With": "XMLHttpRequest",  # Rails UJS sends this
                "Referer": f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}/dropbox",
            },
            follow_redirects=True,
        )

    if response.status_code in (200, 201, 204):
        return {
            "success": True,
            "file": upload_name,
            "size_bytes": file_size,
            "task_url": f"{_base()}/student/classes/{class_id}/core_tasks/{task_id}",
            "http_status": response.status_code,
            "server_response": response.text[:500] if response.text else "",
        }
    else:
        return {
            "success": False,
            "error": f"Upload failed — HTTP {response.status_code}",
            "server_response": response.text[:300],
        }


# ---------------------------------------------------------------------------
# find_task
# ---------------------------------------------------------------------------

async def find_task(query: str) -> dict | None:
    # URL mode — extract class_id and task_id from a ManageBac URL
    url_match = re.search(r'/classes/(\d+)/core_tasks/(\d+)', query)
    if url_match:
        class_id, task_id = url_match.groups()
        return await fetch_task_detail(class_id, task_id)

    # Title search mode — fuzzy match across all classes
    import difflib
    classes = await fetch_classes()
    best_score = 0.0
    best_task = None
    best_class_id = None

    for cls in classes:
        tasks = await fetch_tasks(cls["id"])
        titles = [t["title"] for t in tasks]
        matches = difflib.get_close_matches(query, titles, n=1, cutoff=0.6)
        if matches:
            score = difflib.SequenceMatcher(None, query.lower(), matches[0].lower()).ratio()
            if score > best_score:
                best_score = score
                matched_task = next(t for t in tasks if t["title"] == matches[0])
                best_task = matched_task
                best_class_id = cls["id"]

    if best_task and best_class_id:
        return await fetch_task_detail(best_class_id, best_task["id"])

    return None
